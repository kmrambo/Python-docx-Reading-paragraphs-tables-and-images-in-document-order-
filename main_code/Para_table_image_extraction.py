#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Oct 22 20:59:46 2019

@author: karthick
"""
###Import all necessary packages
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64


#Load the docx file into document object. You can input your own docx file in this step by changing the input path below:
document = Document('/Users/karthick/Desktop/iclouddrive/Work/QA/microchip datasheets/22100F-converted-latest.docx')




##This function extracts the tables and paragraphs from the document object
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

    
#This function extracts the table from the document object as a dataframe
def read_docx_tables(tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

#    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise




combined_df = pd.DataFrame(columns=['para_text','table_id','style'])
table_mod = pd.DataFrame(columns=['string_value','table_id'])
image_df = pd.DataFrame(columns=['image_index','image_rID','image_filename','image_base64_string'])
table_list=[]
xml_list=[]

i=0
imagecounter = 0


blockxmlstring = ''
for block in iter_block_items(document):
    if 'text' in str(block):
        isappend = False
        
        runboldtext = ''
        for run in block.runs:                        
            if run.bold:
                runboldtext = runboldtext + run.text
                
        style = str(block.style.name)
   
        appendtxt = str(block.text)
        appendtxt = appendtxt.replace("\n","")
        appendtxt = appendtxt.replace("\r","")
        tabid = 'Novalue'
        paragraph_split = appendtxt.lower().split()                
        
        isappend = True
        for run in block.runs:
            xmlstr = str(run.element.xml)
            my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
            root = ET.fromstring(xmlstr) 
            #Check if pic is there in the xml of the element. If yes, then extract the image data
            if 'pic:pic' in xmlstr:
                xml_list.append(xmlstr)
                for pic in root.findall('.//pic:pic', my_namespaces):
                    cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                    name_attr = cNvPr_elem.get("name")
                    blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                    embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    isappend = True
                    appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                    document_part = document.part
                    image_part = document_part.related_parts[embed_attr]
                    image_base64 = base64.b64encode(image_part._blob)
                    image_base64 = image_base64.decode()                            
                    dftemp = pd.DataFrame({'image_index':[imagecounter],'image_rID':[embed_attr],'image_filename':[name_attr],'image_base64_string':[image_base64]})
                    image_df = image_df.append(dftemp,sort=False)
                    style = 'Novalue'
                imagecounter = imagecounter + 1
            
    elif 'table' in str(block):
        isappend = True
        style = 'Novalue'
        appendtxt = str(block)
        tabid = i
        dfs = read_docx_tables(tab_id=i)
        dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[i],'style':[style]})
        table_mod = table_mod.append(dftemp,sort=False)
        table_list.append(dfs)
        i=i+1
    if isappend:
            dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[tabid],'style':[style]})
            combined_df=combined_df.append(dftemp,sort=False)
            
combined_df = combined_df.reset_index(drop=True)
image_df = image_df.reset_index(drop=True)
